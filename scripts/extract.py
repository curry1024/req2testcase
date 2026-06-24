"""
需求文档解析脚本
支持格式: PDF (.pdf), Word (.docx)
支持功能: 文本提取、内嵌图片提取

用法:
  python extract.py <file_path>                     # 纯文本提取
  python extract.py <file_path> --extract-images     # 文本 + 内嵌图片提取
  python extract.py <file_path> --extract-images --image-dir <dir>  # 指定图片输出目录
"""
import sys
import os
import zipfile
from pathlib import Path


def extract_pdf(filepath: str) -> str:
    import pdfplumber
    lines = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
    return "\n\n".join(lines)


def extract_docx_text(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            lines.append(row_text)
    return "\n".join(lines)


def extract_docx_images(filepath: str, output_dir: str) -> list:
    """
    从 DOCX (ZIP) 中提取内嵌图片
    DOCX 内图片路径: word/media/image1.png, word/media/image2.jpeg 等
    """
    extracted = []
    os.makedirs(output_dir, exist_ok=True)

    with zipfile.ZipFile(filepath, "r") as zf:
        for name in zf.namelist():
            if name.startswith("word/media/") and not name.endswith("/"):
                basename = Path(name).name
                out_path = os.path.join(output_dir, basename)
                with zf.open(name) as src, open(out_path, "wb") as dst:
                    dst.write(src.read())
                extracted.append(out_path)

    return extracted


def extract_pdf_images(filepath: str, output_dir: str) -> list:
    """
    从 PDF 中提取内嵌图片 (需要 PyMuPDF)
    """
    try:
        import fitz
    except ImportError:
        print("[WARN] PyMuPDF not installed. Install with: pip install PyMuPDF", file=sys.stderr)
        return []

    extracted = []
    os.makedirs(output_dir, exist_ok=True)

    doc = fitz.open(filepath)
    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images(full=True)
        for img_idx, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            ext = base_image["ext"]
            filename = f"page{page_num+1}_img{img_idx+1}.{ext}"
            out_path = os.path.join(output_dir, filename)
            with open(out_path, "wb") as f:
                f.write(image_bytes)
            extracted.append(out_path)

    doc.close()
    return extracted


def get_file_meta(filepath: str) -> dict:
    path = Path(filepath)
    ext = path.suffix.lower()
    size = path.stat().st_size
    return {
        "filename": path.name,
        "format": ext,
        "size_bytes": size,
    }


def main():
    args = sys.argv[1:]
    if not args:
        print("Usage: python extract.py <file_path> [--extract-images] [--image-dir <dir>]", file=sys.stderr)
        sys.exit(1)

    filepath = args[0]
    extract_images_flag = "--extract-images" in args
    image_dir = None

    if "--image-dir" in args:
        idx = args.index("--image-dir")
        if idx + 1 < len(args):
            image_dir = args[idx + 1]

    ext = Path(filepath).suffix.lower()

    if ext not in (".pdf", ".docx"):
        print(f"[ERROR] Unsupported format: {ext}", file=sys.stderr)
        sys.exit(1)

    if not os.path.exists(filepath):
        print(f"[ERROR] File not found: {filepath}", file=sys.stderr)
        sys.exit(1)

    meta = get_file_meta(filepath)
    print(f"[META] file={meta['filename']} format={meta['format']} size={meta['size_bytes']}")

    # 提取文本
    if ext == ".pdf":
        text = extract_pdf(filepath)
    else:  # .docx
        text = extract_docx_text(filepath)

    print(text)

    # 提取内嵌图片
    if extract_images_flag:
        if image_dir is None:
            image_dir = Path(filepath).parent / "extracted_images"
        image_dir = str(image_dir)

        if ext == ".docx":
            image_paths = extract_docx_images(filepath, image_dir)
        else:  # .pdf
            image_paths = extract_pdf_images(filepath, image_dir)

        if image_paths:
            paths_str = ";".join(image_paths)
            print(f"[IMAGES] count={len(image_paths)}")
            print(f"[IMAGE_PATHS] {paths_str}")
        else:
            print("[IMAGES] count=0")


if __name__ == "__main__":
    main()
